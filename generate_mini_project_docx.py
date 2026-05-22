#!/usr/bin/env python3
"""
Generate ES 192 Mini-Project Brief as Word (.docx) for easy editing and Google Docs upload.
"""
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

ARTIFACTS = "/Users/cee-loaners/.gemini/antigravity/brain/c60ff80d-d58b-48f2-9dd7-c3a836701ca0/artifacts"
OUTPUT = "/Users/cee-loaners/Desktop/Projects/ES 192/ES192_Mini_Project_Brief_Skateboard_Deck.docx"

IMG_EXPLODED = os.path.join(ARTIFACTS, "skateboard_exploded_view.png")
IMG_SELECTION = os.path.join(ARTIFACTS, "skateboard_selection_logic.png")
IMG_OPTIONS = os.path.join(ARTIFACTS, "skateboard_material_options.png")
IMG_WEEKLY = os.path.join(ARTIFACTS, "weekly_lab_progression.png")

# Colors
NAVY = RGBColor(25, 60, 120)
DARK = RGBColor(30, 30, 30)
GRAY = RGBColor(100, 100, 100)
WHITE = RGBColor(255, 255, 255)
LIGHT_BLUE_BG = RGBColor(230, 238, 250)


def set_cell_shading(cell, color_hex):
    """Set background color of a table cell."""
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color_hex)
    shading.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(shading)


def add_styled_table(doc, headers, rows, col_widths=None):
    """Add a professional-looking table."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'

    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = WHITE
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, '193C78')

    # Data rows
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = ''
            p = cell.paragraphs[0]
            run = p.add_run(val)
            run.font.size = Pt(9)
            run.font.color.rgb = DARK
            if ri % 2 == 0:
                set_cell_shading(cell, 'EDF2FA')

    # Set column widths if provided
    if col_widths:
        for ri, row in enumerate(table.rows):
            for ci, w in enumerate(col_widths):
                row.cells[ci].width = Inches(w)

    doc.add_paragraph()  # spacing
    return table


def add_heading_styled(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = NAVY
    return h


def add_bullet(doc, text, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        run.font.size = Pt(10)
        p.add_run(text).font.size = Pt(10)
    else:
        run = p.add_run(text)
        run.font.size = Pt(10)
    return p


def add_body(doc, text):
    p = doc.add_paragraph(text)
    for run in p.runs:
        run.font.size = Pt(10)
    return p


def add_note_box(doc, title, text):
    """Add a highlighted note/callout."""
    p = doc.add_paragraph()
    run = p.add_run(f'{title}: ')
    run.bold = True
    run.font.size = Pt(9)
    run.font.color.rgb = NAVY
    run2 = p.add_run(text)
    run2.font.size = Pt(9)
    run2.font.color.rgb = RGBColor(60, 60, 60)
    run2.italic = True
    return p


def add_image_centered(doc, path, width=5.5, caption=None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(path, width=Inches(width))
    if caption:
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_cap = cap.add_run(caption)
        run_cap.italic = True
        run_cap.font.size = Pt(8)
        run_cap.font.color.rgb = GRAY


def build_docx():
    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(10)
    font.color.rgb = DARK

    # Adjust margins
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # =========================================================================
    # TITLE PAGE
    # =========================================================================
    for _ in range(4):
        doc.add_paragraph()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('ES 192: Mini-Project Brief')
    run.bold = True
    run.font.size = Pt(28)
    run.font.color.rgb = NAVY

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run('The Multi-Material Skateboard Deck Challenge')
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(60, 60, 60)

    doc.add_paragraph()  # spacing

    # Horizontal line
    line_p = doc.add_paragraph()
    line_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = line_p.add_run('_' * 60)
    run.font.color.rgb = NAVY
    run.font.size = Pt(10)

    doc.add_paragraph()

    for text in [
        'Materials Selection in Mechanical Design',
        'Harvard SEAS - Fall 2026',
        '',
        'Prepared by: Seymur Hasanov',
        'For discussion with: Nora Cullen, Director for Active Learning',
        '',
        'Draft - April 2026',
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        run.font.size = Pt(11)
        run.font.color.rgb = GRAY
        if 'Draft' in text:
            run.italic = True

    doc.add_paragraph()
    add_image_centered(doc, IMG_EXPLODED, width=4.0)

    doc.add_page_break()

    # =========================================================================
    # SECTION 1 - CONCEPT OVERVIEW
    # =========================================================================
    add_heading_styled(doc, '1. Concept Overview', level=1)

    add_body(doc,
        'Each team of 3 students designs, fabricates, and tests a multi-material mini skateboard deck '
        'that must maximize load-carrying capacity while minimizing weight. The deck is a sandwich panel: '
        'laser-cut core layers bonded to cast composite skins, with 3D-printed mounting hardware.')

    add_body(doc,
        'What makes this a materials selection project (not just a build project): every component\'s '
        'material is chosen using the Ashby methodology - students derive materials indices, screen '
        'candidates on property charts, and justify their choices with measured data. They then validate '
        'predictions against physical test results.')

    add_heading_styled(doc, 'Why a Skateboard Deck?', level=2)
    add_bullet(doc, 'It is fundamentally a beam/panel under bending - same mechanics as a cantilever, but far more engaging')
    add_bullet(doc, 'The sandwich structure (skins + core) directly teaches composite and hybrid material concepts')
    add_bullet(doc, '3-point bend testing is standard and requires a simpler fixture than a cantilever clamp')
    add_bullet(doc, 'Students can hold it, flex it, and keep it - tangible and rewarding')
    add_bullet(doc, 'Directly leads into the "Composite longboard deck" group project option already in the syllabus')

    add_heading_styled(doc, 'Performance Metric', level=2)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('P  =  F_max  /  m')
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = NAVY

    add_bullet(doc, 'F_max = maximum load supported without exceeding 5 mm center deflection')
    add_bullet(doc, 'm = total mass of the deck assembly')
    add_bullet(doc, 'Secondary metric: load at structural failure / mass')

    # =========================================================================
    # SECTION 2 - SPECIFICATIONS
    # =========================================================================
    add_heading_styled(doc, '2. Specifications', level=1)

    add_styled_table(doc,
        ['Parameter', 'Value'],
        [
            ['Deck length', '~300 mm (mini/cruiser scale)'],
            ['Deck width', '~80 mm'],
            ['Total thickness', '~12-15 mm (core + skins)'],
            ['Shape', 'Flat (no concave, no kick tails)'],
            ['Test method', '3-point bend: supports at truck positions (~200 mm span), load at center'],
            ['Max allowable deflection', '5 mm under service load'],
            ['Performance metric', 'P = F_max / m'],
            ['Manufacturing processes', '>= 3 required: laser cutting, 3D printing, silicone mold casting'],
            ['Materials', 'Standard ALL stock provided (see Section 5)'],
        ],
        col_widths=[2.0, 4.5]
    )

    doc.add_page_break()

    # =========================================================================
    # SECTION 3 - STRUCTURAL CONCEPT
    # =========================================================================
    add_heading_styled(doc, '3. Structural Concept - Exploded View', level=1)

    add_body(doc,
        'The deck is a sandwich panel - the same architecture used in real performance skateboard decks, '
        'aerospace panels, and composite structures. Each layer serves a specific structural function and '
        'is made by a different manufacturing process.')

    add_image_centered(doc, IMG_EXPLODED, width=5.5,
        caption='Figure 1: Exploded isometric view of the multi-material skateboard deck. '
                'Laser-cut plywood core (blue), cast fiber-reinforced epoxy skins (green), '
                '3D-printed hardware (orange). Bottom: assembled profile with 3-point bend test.')

    add_heading_styled(doc, 'Sandwich Panel Mechanics', level=2)
    add_bullet(doc, 'carry tension and compression from bending', bold_prefix='Top & bottom skins (cast composite): ')
    add_bullet(doc, 'carry shear, provide bulk thickness to separate skins', bold_prefix='Core layers (laser-cut plywood): ')
    add_bullet(doc, 'truck mounts transfer load, bumpers absorb impact', bold_prefix='Hardware (3D-printed): ')

    add_body(doc,
        'This sandwich architecture is a direct application of the course\'s hybrid materials content '
        '(Chapters 13-14). Students learn that separating stiff skins with a lightweight core dramatically '
        'increases bending stiffness with minimal weight penalty - a concept they can feel by flexing their deck.')

    doc.add_page_break()

    # =========================================================================
    # SECTION 4 - MATERIALS SELECTION LOGIC
    # =========================================================================
    add_heading_styled(doc, '4. Materials Selection Logic', level=1)

    add_body(doc,
        'The Ashby methodology drives the material-to-process assignment for each component. '
        'Each component has a function, which translates to a materials index, which points to '
        'candidate materials, which constrain the manufacturing process.')

    add_image_centered(doc, IMG_SELECTION, width=5.8,
        caption='Figure 2: Materials selection logic - Function -> Index -> Ashby Chart -> Material -> Process. '
                'Bottom: cross-section showing sandwich panel mechanics.')

    doc.add_page_break()

    # =========================================================================
    # SECTION 5 - MATERIAL OPTIONS
    # =========================================================================
    add_heading_styled(doc, '5. Material Options - The Decision Space', level=1)

    add_body(doc,
        'The whole point of the Ashby exercise is that students have multiple real candidates to '
        'choose from and must justify their selection with data. Below is the full palette of materials '
        'available in the ALL, organized by manufacturing process.')

    add_image_centered(doc, IMG_OPTIONS, width=5.8,
        caption='Figure 3: Complete material candidate pool organized by manufacturing process, '
                'with key properties and selection questions.')

    doc.add_page_break()

    # 5.1 Core
    add_heading_styled(doc, '5.1  Core Layers - Laser Cut (students pick 1-2)', level=2)
    add_styled_table(doc,
        ['Material', 'E (GPa)', 'Density (kg/m3)', 'E^(1/2)/rho', 'Notes'],
        [
            ['Birch plywood (3mm)', '12', '600', '5.8 x 10^-3', 'Natural fiber composite, best index'],
            ['Acrylic / PMMA (3mm)', '3.2', '1180', '1.5 x 10^-3', 'Transparent, brittle, poor index'],
            ['MDF (3mm)', '4', '750', '2.7 x 10^-3', 'Uniform (no grain), heavy'],
            ['Aluminum (1.5mm)', '69', '2700', '3.1 x 10^-3', 'Highest E, but heaviest'],
        ],
        col_widths=[1.3, 0.6, 0.9, 0.8, 2.9]
    )
    add_note_box(doc, 'Selection question',
        'Plywood wins on E^(1/2)/rho, but aluminum has 6x higher absolute E. '
        'When does raw stiffness matter more than specific stiffness? Students test coupons in Week 2 to decide.')

    # 5.2 Skins
    add_heading_styled(doc, '5.2  Composite Skins - Mold Cast (students pick 1 per skin)', level=2)
    add_styled_table(doc,
        ['Material', 'E (GPa)', 'Density (kg/m3)', 'Notes'],
        [
            ['Neat epoxy', '~3', '~1200', 'Baseline, no filler, easy to pour'],
            ['Rigid polyurethane', '~2.5', '~1100', 'Fast cure (~30 min), lower cost'],
            ['Chopped glass + epoxy', '~8-12', '~1600', 'Short fiber composite, highest stiffness'],
            ['Mineral-filled epoxy', '~5-6', '~1500', 'Stiffer than neat, cheaper than glass'],
            ['Flexible polyurethane', '~0.01', '~1050', 'NOT structural - grip/damping layer'],
        ],
        col_widths=[1.5, 0.6, 0.9, 3.5]
    )
    add_note_box(doc, 'Selection question',
        'Max stiffness (glass-filled bottom skin) or trade top-skin stiffness for a flexible grip layer? '
        'What if you use glass-filled on the bottom (tension side) and flexible PU on top (grip)?')

    # 5.3 Hardware
    add_heading_styled(doc, '5.3  Hardware - 3D Printed (students pick per component)', level=2)
    add_styled_table(doc,
        ['Material', 'E (GPa)', 'Strength (MPa)', 'Notes'],
        [
            ['PLA', '3.5', '60', 'Stiffest, easiest to print, brittle - cracks on impact'],
            ['PETG', '2.2', '50', 'Tougher, better layer adhesion, slight flex'],
            ['TPU', '0.03', '30', 'Flexible, impact absorbing - bumpers only?'],
        ],
        col_widths=[1.0, 0.6, 0.8, 4.1]
    )
    add_note_box(doc, 'Selection question',
        'Truck mounts see bolt preload + vibration. PLA is stiffer but PETG won\'t crack. '
        'Students justify using sigma_f/rho AND consider failure mode (brittle vs. yielding).')

    doc.add_page_break()

    # =========================================================================
    # SECTION 6 - LECTURE-LAB ALIGNMENT
    # =========================================================================
    add_heading_styled(doc, '6. Lecture-Lab Alignment', level=1)

    add_body(doc,
        'A key design principle of this mini-project is that each lab aligns with the lecture content '
        'being taught that same week. The table below shows how the weekly lab activities map to the '
        'textbook chapters covered in parallel lectures.')

    add_styled_table(doc,
        ['Week', 'Lecture Topic (Chapters)', 'Lab Activity', 'How They Align'],
        [
            ['1', 'Ch 1-2: Intro, material properties, material families',
             'Design + Ashby charts',
             'Students apply Ch 2 property definitions and Ch 1 material families to their deck design'],
            ['2', 'Ch 3-4: Ashby charts, selection strategy, materials indices',
             'Laser cutting + 3-point bend tests',
             'Students generate Ashby charts and use indices to justify their core material - direct application of Ch 3-4'],
            ['3', 'Ch 5: Case studies applying the methodology',
             '3D printing + orientation experiment',
             'Case studies show real selection decisions; lab shows how process affects properties'],
            ['4', 'Ch 6: Processes and their effect on properties',
             'Silicone mold casting + resin selection',
             'Ch 6 covers process-property relationships; students directly observe how casting and fiber addition change E'],
            ['5', 'Ch 7-8: Process selection, AM, design for AM',
             'Load testing + validation',
             'Students reflect on all three process choices (Ch 7) and validate their AM decisions (Ch 8)'],
            ['6', 'Ch 9-10: Multiple constraints, trade-offs',
             'COMSOL simulation',
             'Simulation enables "what-if" exploration of trade-offs (Ch 9-10)'],
        ],
        col_widths=[0.4, 1.7, 1.5, 2.9]
    )

    add_heading_styled(doc, 'Note on Composites Content', level=2)
    add_body(doc,
        'The formal composites theory (rule of mixtures, bounds, sandwich panel analysis) is covered '
        'in Weeks 8-9 lectures (Chapters 13-14). However, in Week 4, students are casting fiber-reinforced '
        'resin skins. This is intentional and pedagogically sound for the following reasons:')

    add_bullet(doc, 'Week 4 lectures cover Chapter 6: "How processing changes microstructure and properties." '
        'Adding chopped fiber to resin is a direct example of a process-property relationship, '
        'which is exactly the Ch 6 content.')
    add_bullet(doc, 'Students do NOT need rule-of-mixtures theory to select a resin system. They need to: '
        '(a) observe that fiber-filled resin has higher E than neat resin, (b) place it on an Ashby chart, '
        'and (c) justify the choice. This is pure Ashby methodology (Ch 3-4), applied to a process outcome (Ch 6).')
    add_bullet(doc, 'The experiential introduction in Week 4 ("I cast this, it\'s stiffer") creates a foundation '
        'that the formal composite theory in Weeks 8-9 can build on. Students will already know what a '
        'fiber-reinforced material feels like before they learn the math. This is experience-first pedagogy.')
    add_bullet(doc, 'The group project (Weeks 7-13) is where students apply the full composite theory '
        'from Ch 13-14 - the "Composite longboard deck" project is explicitly designed for this.')

    add_note_box(doc, 'Summary',
        'Week 4 lab teaches composites as a process-property concept (Ch 6). '
        'Weeks 8-9 lectures formalize it as composite theory (Ch 13-14). '
        'The group project applies the full theory. This sequence: experience -> theory -> application.')

    doc.add_page_break()

    # =========================================================================
    # SECTION 7 - LAB SCHEDULE OVERVIEW
    # =========================================================================
    add_heading_styled(doc, '7. Lab Schedule Overview (6 Labs)', level=1)

    add_body(doc,
        'The mini-project runs over 6 weekly labs (2 hours each, ~10 students / 3-4 teams per section). '
        'Each lab has two parallel tracks: a manufacturing skill and a materials selection activity. '
        'The selection thread builds cumulatively: Screen -> Measure -> Validate -> Select -> Compare -> Simulate.')

    add_image_centered(doc, IMG_WEEKLY, width=6.0,
        caption='Figure 4: Six-week lab progression showing manufacturing skills (top) and '
                'materials selection activities (bottom).')

    doc.add_page_break()

    # =========================================================================
    # SECTION 8 - DETAILED LAB PLANS
    # =========================================================================
    add_heading_styled(doc, '8. Detailed Lab Plans', level=1)

    # --- LAB 1 ---
    add_heading_styled(doc, 'Lab 1 - Design & Planning (Week 1)', level=2)

    p = doc.add_paragraph()
    run = p.add_run('Selection Concept: SCREEN')
    run.bold = True
    run.font.color.rgb = NAVY
    run.font.size = Pt(10)

    add_heading_styled(doc, 'Objectives', level=3)
    add_bullet(doc, 'Students are safe to work in the ALL and understand all equipment')
    add_bullet(doc, 'Teams of 3 form and understand the skateboard deck challenge')
    add_bullet(doc, 'Each team produces a design plan with Ashby-based material justification')

    add_heading_styled(doc, 'Plan (2 hours)', level=3)
    add_styled_table(doc,
        ['Time', 'Activity', 'Details'],
        [
            ['0:00-0:30', 'ALL Tour + Safety', 'Walk through laser bay, 3D print lab, casting station, workshop. Safety rules. Sign agreements.'],
            ['0:30-0:45', 'Project Briefing', 'Present the Skateboard Deck Challenge: specs, metric, test setup, timeline, grading rubric.'],
            ['0:45-1:00', 'Team Formation', 'Form teams of 3. Concept sketching: shape, layers, where do skins go?'],
            ['1:00-1:30', 'Design + Selection', 'Design Requirements Tables per component. Ashby E vs rho chart. Derive index. Screen candidates.'],
            ['1:30-1:50', 'Design Decisions', 'Assign material + process per component. Begin rough dimensioning. CAs challenge choices.'],
            ['1:50-2:00', 'Wrap-up', 'Deliverable explained: Design Brief (1-2 pages) due before next lab.'],
        ],
        col_widths=[0.7, 1.1, 4.7]
    )

    add_heading_styled(doc, 'Materials & Equipment', level=3)
    add_bullet(doc, 'Whiteboards or large paper, markers')
    add_bullet(doc, 'Printed material property data sheets (E, rho, sigma_f for all stock)')
    add_bullet(doc, 'Rulers, example deck photos, materials selection software access')

    p = doc.add_paragraph()
    run = p.add_run('Deliverable: ')
    run.bold = True
    run.font.size = Pt(10)
    p.add_run('Design Brief (1-2 pages): annotated sketch, Design Requirements Table per component, '
              'Ashby chart with index line and candidates highlighted, material/process justification.').font.size = Pt(10)

    doc.add_page_break()

    # --- LAB 2 ---
    add_heading_styled(doc, 'Lab 2 - Laser Cutting: Core Layers (Week 2)', level=2)

    p = doc.add_paragraph()
    run = p.add_run('Selection Concept: MEASURE')
    run.bold = True
    run.font.color.rgb = NAVY

    add_heading_styled(doc, 'Objectives', level=3)
    add_bullet(doc, 'Students can safely operate the laser cutter')
    add_bullet(doc, 'Core layers are cut and test coupons prepared from 2+ materials')
    add_bullet(doc, 'Students have measured flexural modulus to validate Ashby chart data')

    add_heading_styled(doc, 'Plan (2 hours)', level=3)
    add_styled_table(doc,
        ['Time', 'Activity', 'Details'],
        [
            ['0:00-0:20', 'Laser Cutter Training', 'Demo: operation, material loading, focus, DXF/SVG import, kerf compensation, speed/power settings. Safety.'],
            ['0:20-0:40', 'File Preparation', 'Teams prep cut files: core layer profiles + test coupons (100x20mm from 2+ materials).'],
            ['0:40-1:10', 'Cutting', 'Teams take turns. Stock: 3mm plywood, 3mm acrylic, 3mm MDF, 1.5mm aluminum.'],
            ['1:10-1:40', 'Baseline Testing', '3-point bend test: 80mm span, load at center, measure deflection at 3+ loads. Calculate E = FL^3/(48*delta*I).'],
            ['1:40-1:55', 'Discussion', 'Which material had best specific stiffness? Did measurements match database? Process constraints?'],
            ['1:55-2:00', 'Wrap-up', 'Verify cut core layers + test data. Remind: bring CAD files for 3D printing next week.'],
        ],
        col_widths=[0.7, 1.1, 4.7]
    )

    add_heading_styled(doc, 'Materials & Equipment', level=3)
    add_bullet(doc, 'Sheet stock: 3mm birch plywood, 3mm acrylic, 3mm MDF, 1.5mm aluminum')
    add_bullet(doc, '3-point bend test blocks (2 support + loading nose), calibrated weights (100g-1kg)')
    add_bullet(doc, 'Dial indicator or rulers, digital scales, calipers')

    p = doc.add_paragraph()
    run = p.add_run('Deliverable: ')
    run.bold = True
    p.add_run('Laser-cut core layers + measured flexural modulus data for 2+ materials with comparison to Ashby database.')

    doc.add_page_break()

    # --- LAB 3 ---
    add_heading_styled(doc, 'Lab 3 - 3D Printing: Hardware & Mold Pattern (Week 3)', level=2)

    p = doc.add_paragraph()
    run = p.add_run('Selection Concept: VALIDATE PROCESS-PROPERTY EFFECTS')
    run.bold = True
    run.font.color.rgb = NAVY

    add_heading_styled(doc, 'Objectives', level=3)
    add_bullet(doc, 'Students can use FDM printers and slicer software')
    add_bullet(doc, 'Truck mounts, bumpers, and mold pattern are printing')
    add_bullet(doc, 'Students understand print orientation anisotropy through hands-on experiment')

    add_heading_styled(doc, 'Plan (2 hours)', level=3)
    add_styled_table(doc,
        ['Time', 'Activity', 'Details'],
        [
            ['0:00-0:20', '3D Printing Training', 'Demo: slicer (STL, layer height, infill, orientation, supports). Filaments: PLA, PETG, TPU.'],
            ['0:20-0:35', 'Design Review + Slicing', 'CAs review printability. Prep: (A) truck mounts/bumpers, (B) mold pattern with draft angles.'],
            ['0:35-0:55', 'Orientation Experiment', 'Print 2 test beams (60x10x5mm) in PLA: flat vs. on-edge (~10 min each).'],
            ['0:55-1:15', 'Test Orientation Beams', '3-point bend both. Calculate stiffness ratio. Flat beam is stiffer - demonstrates anisotropy.'],
            ['1:15-1:45', 'Start Functional Prints', 'Queue truck mounts, bumpers, mold pattern. Small parts finish in session.'],
            ['1:45-2:00', 'Wrap-up', 'Confirm mold patterns ready for Week 4. Preview casting. Bring ALL components next week.'],
        ],
        col_widths=[0.7, 1.1, 4.7]
    )

    add_heading_styled(doc, 'Materials & Equipment', level=3)
    add_bullet(doc, 'PLA, PETG, TPU filament spools; FDM printers (2-3 per section)')
    add_bullet(doc, 'Slicer software on lab computers; pre-made test beam STL (via Canvas)')
    add_bullet(doc, '3-point bend test setup from Week 2')

    p = doc.add_paragraph()
    run = p.add_run('Deliverable: ')
    run.bold = True
    p.add_run('3D-printed truck mounts + bumpers + mold pattern + print orientation anisotropy data.')

    doc.add_page_break()

    # --- LAB 4 ---
    add_heading_styled(doc, 'Lab 4 - Silicone Mold Casting & Assembly (Week 4)', level=2)

    p = doc.add_paragraph()
    run = p.add_run('Selection Concept: CREATE NEW MATERIALS')
    run.bold = True
    run.font.color.rgb = NAVY

    add_heading_styled(doc, 'Objectives', level=3)
    add_bullet(doc, 'Students can make silicone molds and cast resin parts')
    add_bullet(doc, 'Composite skin layers are cast with justified resin selection')
    add_bullet(doc, 'Deck assembly begins - bonding skins to core')

    add_heading_styled(doc, 'Plan (2 hours)', level=3)
    add_styled_table(doc,
        ['Time', 'Activity', 'Details'],
        [
            ['0:00-0:15', 'Casting Training', 'Demo: mold release, containment box (foam core + hot glue), mix silicone by weight, pour, degas. Cure: 4-6 hrs.'],
            ['0:15-0:40', 'Mold-Making', 'Pour silicone molds for flat skin panels. Already-cured teams proceed to resin casting. Backup molds available.'],
            ['0:40-1:10', 'Resin Casting', 'Select resin (must justify): neat epoxy, rigid PU, fiber-filled epoxy, mineral-filled, or flexible PU. Mix, add fibers if chosen, pour.'],
            ['1:10-1:40', 'Begin Assembly', 'Dry-fit core layers + truck mounts. Plan bonding (epoxy, bolts, press-fit). Begin laminating core layers.'],
            ['1:40-1:55', 'Discussion', 'What shapes does casting enable? Defects observed? How does resin choice affect performance?'],
            ['1:55-2:00', 'Wrap-up', 'CAs demold after session. Teams finalize assembly before Week 5.'],
        ],
        col_widths=[0.7, 1.1, 4.7]
    )

    add_heading_styled(doc, 'Materials & Equipment', level=3)
    add_bullet(doc, 'Silicone rubber (two-part, e.g., Smooth-On Mold Star), mold release spray')
    add_bullet(doc, 'Foam core sheets + hot glue guns (containment boxes)')
    add_bullet(doc, 'Casting resins: rigid PU, flexible PU, epoxy, chopped glass fiber, mineral filler')
    add_bullet(doc, 'Mixing cups, stir sticks, digital scale, gloves, eye protection')
    add_bullet(doc, 'Vacuum degassing chamber (if available); pre-made backup molds (2-3 per section)')

    p = doc.add_paragraph()
    run = p.add_run('Deliverable: ')
    run.bold = True
    p.add_run('Cast composite skin layers + resin selection justification + partially assembled deck.')

    doc.add_page_break()

    # --- LAB 5 ---
    add_heading_styled(doc, 'Lab 5 - Load Testing & Validation (Week 5)', level=2)

    p = doc.add_paragraph()
    run = p.add_run('Selection Concept: COMPARE PREDICTED vs. ACTUAL')
    run.bold = True
    run.font.color.rgb = NAVY

    add_heading_styled(doc, 'Objectives', level=3)
    add_bullet(doc, 'Every team has a fully assembled, tested deck with load-deflection data')
    add_bullet(doc, 'Students compare Ashby predictions to actual performance')
    add_bullet(doc, 'Students analyze failure modes and reflect on methodology gaps')

    add_heading_styled(doc, 'Plan (2 hours)', level=3)
    add_styled_table(doc,
        ['Time', 'Activity', 'Details'],
        [
            ['0:00-0:15', 'Final Assembly', 'Attach skins to core. Mount truck plates. Trim flash. Weigh. Measure dimensions.'],
            ['0:15-0:20', 'Test Procedure', 'Support at truck positions (~200mm span). Load center incrementally. Record F vs delta. Find F_max at 5mm.'],
            ['0:20-1:00', 'Testing', '~10 min per team. Others: calculate P, review Week 1 predictions. Load to failure if safe.'],
            ['1:00-1:20', 'Leaderboard', 'Calculate P = F_max/m. Post results. Rank by performance. Celebrate top design.'],
            ['1:20-1:50', 'Validation Discussion', 'Ashby prediction vs actual. Why different? Which choices mattered? What would you change?'],
            ['1:50-2:00', 'Report Assignment', 'Report (3-5 pages): design, Ashby plots, test data, predicted vs actual, failure modes, lessons.'],
        ],
        col_widths=[0.7, 1.1, 4.7]
    )

    add_heading_styled(doc, 'Materials & Equipment', level=3)
    add_bullet(doc, '3-point bend fixture (2 supports at ~200mm span), calibrated weights (100g-5kg)')
    add_bullet(doc, 'Dial indicator + magnetic stand, digital scale, calipers')
    add_bullet(doc, 'Whiteboard for leaderboard, camera for failure photos')

    p = doc.add_paragraph()
    run = p.add_run('Deliverable: ')
    run.bold = True
    p.add_run('Report (3-5 pages, due in 1 week): design rationale, Ashby plots, load-deflection curve, '
              'predicted vs actual, failure mode analysis, lessons learned.')

    doc.add_page_break()

    # --- LAB 6 ---
    add_heading_styled(doc, 'Lab 6 - COMSOL Simulation (Week 6)', level=2)

    p = doc.add_paragraph()
    run = p.add_run('Selection Concept: SIMULATE & THREE-WAY VALIDATION')
    run.bold = True
    run.font.color.rgb = NAVY

    add_heading_styled(doc, 'Objectives', level=3)
    add_bullet(doc, 'Students can set up a basic structural simulation in COMSOL')
    add_bullet(doc, 'Three-way comparison: Ashby prediction vs. COMSOL vs. experiment')
    add_bullet(doc, 'Students understand when simulation adds value vs. when testing is necessary')

    add_heading_styled(doc, 'Plan (2 hours) - Run by COMSOL Instruction Team', level=3)
    add_styled_table(doc,
        ['Time', 'Activity', 'Details'],
        [
            ['0:00-0:20', 'COMSOL Intro', 'Interface overview. Pre-built sandwich panel template. Geometry, materials, BCs, load.'],
            ['0:20-0:50', 'Guided Exercise', 'Modify template: deck dimensions, assign measured material properties, apply test loads. Run.'],
            ['0:50-1:10', 'Comparison Exercise', 'Compare COMSOL deflection to Week 5 data. Fill comparison table for each load level.'],
            ['1:10-1:40', 'Parametric Study', 'Vary one parameter (core material, skin thickness). Plot results. "What if aluminum core?"'],
            ['1:40-2:00', 'Wrap-up', 'When is simulation useful vs testing? Preview COMSOL features for group projects.'],
        ],
        col_widths=[0.7, 1.1, 4.7]
    )

    add_heading_styled(doc, 'Three-Way Comparison (students fill in)', level=3)
    add_styled_table(doc,
        ['Method', 'Source', 'Deflection at Load X'],
        [
            ['Ashby index (Week 1)', 'Analytical, from materials index', 'delta_predicted'],
            ['COMSOL simulation (Week 6)', 'Computational, idealized geometry', 'delta_simulated'],
            ['Experiment (Week 5)', 'Physical test, real structure', 'delta_measured'],
        ],
        col_widths=[1.8, 2.2, 2.5]
    )

    add_heading_styled(doc, 'Materials & Equipment', level=3)
    add_bullet(doc, 'COMSOL via HUIT VDI (access set up in advance)')
    add_bullet(doc, 'Pre-built sandwich panel template (.mph file)')
    add_bullet(doc, 'Material property table, step-by-step worksheet')

    doc.add_page_break()

    # =========================================================================
    # SECTION 9 - SELECTION THREAD SUMMARY
    # =========================================================================
    add_heading_styled(doc, '9. Materials Selection Thread Summary', level=1)

    add_body(doc,
        'The materials selection methodology is reinforced every week. Each lab has a "selection verb" '
        'that builds on the previous week:')

    add_styled_table(doc,
        ['Week', 'Lab', 'Selection Verb', 'What Students Learn'],
        [
            ['1', 'Design', 'SCREEN', 'Use Ashby charts + indices to narrow candidates'],
            ['2', 'Laser Cut', 'MEASURE', 'Real properties != database; validate assumptions'],
            ['3', '3D Print', 'VALIDATE', 'Process changes properties (anisotropy)'],
            ['4', 'Casting', 'CREATE', 'Design new materials (fiber composites) for new chart regions'],
            ['5', 'Testing', 'COMPARE', 'Prediction vs reality; where methodology works/breaks'],
            ['6', 'COMSOL', 'SIMULATE', 'Computational "what-if" exploration'],
        ],
        col_widths=[0.4, 0.8, 1.0, 4.3]
    )

    # =========================================================================
    # SECTION 10 - GRADING
    # =========================================================================
    add_heading_styled(doc, '10. Grading Rubric (10% of Course Grade)', level=1)

    add_styled_table(doc,
        ['Criterion', 'Weight', 'Description'],
        [
            ['Materials Selection Rationale', '30%', 'Ashby charts, indices, Design Requirements Tables, justification quality'],
            ['Design Quality & Manufacturing', '25%', 'Creativity, execution quality, appropriate use of each process'],
            ['Test Results & Performance', '20%', 'Data quality, P = F_max/m, testing rigor'],
            ['Report & Analysis', '25%', 'Clarity, predicted-vs-actual, failure modes, reflection'],
        ],
        col_widths=[1.8, 0.6, 4.1]
    )

    # =========================================================================
    # SECTION 11 - ACTION ITEMS
    # =========================================================================
    add_heading_styled(doc, '11. Pre-Semester Action Items', level=1)

    add_styled_table(doc,
        ['Task', 'Owner', 'Target'],
        [
            ['Finalize deck specs and 3-point bend fixture design', 'Seymur', 'End of April 2026'],
            ['Test-run full 6-week sequence; write lab handouts + CA instructions', 'Nora', 'End of May 2026'],
            ['Build/procure 3-point bend test fixture (2 supports + weights)', 'Nora + ALL staff', 'End of May 2026'],
            ['Prepare 2-3 backup silicone molds (flat skin panel shape)', 'Nora', 'End of May 2026'],
            ['Coordinate with COMSOL team: sandwich panel template specs', 'Seymur', 'End of April 2026'],
            ['COMSOL team: build sandwich panel template .mph file', 'COMSOL team', 'End of May 2026'],
            ['Recruit and train lab CAs from ES 51', 'Seymur + Nora', 'August 2026'],
            ['Procure silicone, resins, chopped fiber, casting supplies', 'Nora + ALL staff', 'August 2026'],
            ['Prepare material data sheets for students (E, sigma_f, rho)', 'Seymur', 'Before Week 1'],
            ['Set up materials selection software + COMSOL VDI access', 'Seymur', 'Before Week 1'],
        ],
        col_widths=[3.5, 1.5, 1.5]
    )

    # =========================================================================
    # SAVE
    # =========================================================================
    doc.save(OUTPUT)
    print(f"Word document saved to: {OUTPUT}")
    print(f"Ready for Google Docs upload or editing in Word.")


if __name__ == "__main__":
    build_docx()
